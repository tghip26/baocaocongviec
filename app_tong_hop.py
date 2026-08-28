import os
import re
import csv
import glob
import sys
import shutil
import unicodedata
import threading
from datetime import datetime

import openpyxl
import docx
from docx import Document

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Đảm bảo UTF-8 cho console
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

# ============================================================
# CẤU HÌNH & CHUẨN HÓA DÙNG CHUNG
# ============================================================
DEFAULT_ORG1 = "ỦY BAN NHÂN DÂN TỈNH BẮC NINH"
DEFAULT_ORG2 = "SỞ Y TẾ"
DEFAULT_ORG3 = "BỆNH VIỆN ĐA KHOA BẮC NINH SỐ 2"
DEFAULT_PROVINCE = "Bắc Ninh"

def clean_text(value):
    if value is None:
        return ""
    text = str(value).replace("\xa0", " ").replace("\r", " ").replace("\n", " ").strip()
    return re.sub(r"\s+", " ", text)

def remove_accents(text):
    if not text:
        return ""
    text = str(text).replace("đ", "d").replace("Đ", "D")
    text = unicodedata.normalize("NFD", text)
    text = "".join(c for c in text if unicodedata.category(c) != "Mn")
    return text.lower()

def norm_name(value):
    text = clean_text(value)
    text = str(text).replace("đ", "d").replace("Đ", "D")
    text = unicodedata.normalize("NFD", text)
    text = "".join(c for c in text if unicodedata.category(c) != "Mn")
    return re.sub(r"[^a-zA-Z0-9]", "", text).lower()

def norm_cccd(value):
    if not value:
        return ""
    digits = re.sub(r"\D", "", str(value))
    return digits.lstrip("0")

def format_cccd_output(value):
    digits = re.sub(r"\D", "", str(value))
    if len(digits) == 11:
        return "0" + digits
    return digits

def normalize_date(value):
    val = clean_text(value)
    if not val:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    match = re.search(r"(\d{1,2})[\/\-.](\d{1,2})[\/\-.](\d{2,4})", val)
    if match:
        day = int(match.group(1))
        month = int(match.group(2))
        year = int(match.group(3))
        if year < 100:
            year += 2000 if year < 50 else 1900
        try:
            return datetime(year, month, day).strftime("%d/%m/%Y")
        except ValueError:
            return val
    return val

def normalize_issue_place(value):
    val = clean_text(value).upper()
    if "CCSĐKQLCTVDLQGVDC" in val or "ĐKQL" in val or "DÂN CƯ" in val or "BỘ CÔNG AN" in val:
        return "CCSĐKQLCTVDLQGVDC"
    return "CCSQLHCVTTXH"

def clean_email(email_str):
    if not email_str:
        return ""
    email_clean = clean_text(email_str)
    if "@" in email_clean:
        parts = email_clean.split("@")
        return parts[0] + "@" + parts[1]
    return email_clean

# ============================================================
# LOGIC TÁC VỤ 1: ĐỐI CHIẾU DỮ LIỆU & IMPORT VGCA (xu_ly_doi_chieu)
# ============================================================
def run_task_doi_chieu(folder_path, log_func, progress_func):
    log_func("=== BẮT ĐẦU TÁC VỤ 1: ĐỐI CHIẾU DỮ LIỆU WORD & EXCEL SSO ===")
    
    # 1. Đọc SSO Excel
    sso_map = {}
    excel_files = glob.glob(os.path.join(folder_path, "*.xlsx"))
    
    for f in excel_files:
        filename = os.path.basename(f)
        if filename.startswith("~$") or filename in ["Ket_qua.xlsx", "Ket_qua_moi.xlsx", "Ket_qua_final.xlsx", "Template_import VGCA.xlsx"]:
            continue
        log_func(f"Đang đọc dữ liệu SSO từ file: {filename}")
        try:
            wb = openpyxl.load_workbook(f, data_only=True)
            for sheetname in wb.sheetnames:
                ws = wb[sheetname]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue
                header = [clean_text(c) for c in rows[0]]
                name_idx, email_idx, cccd_idx = -1, -1, -1
                for idx, col in enumerate(header):
                    c_lower = col.lower()
                    if "tên" in c_lower or "họ tên" in c_lower:
                        if name_idx == -1: name_idx = idx
                    if "tên đăng nhập" in c_lower or "email" in c_lower or "thư điện tử" in c_lower:
                        if email_idx == -1: email_idx = idx
                    if "description" in c_lower or "cccd" in c_lower or "cmt" in c_lower or "cmnd" in c_lower:
                        if cccd_idx == -1: cccd_idx = idx
                
                if name_idx == -1: name_idx = 0
                if email_idx == -1: email_idx = 1
                if cccd_idx == -1 and len(header) >= 6: cccd_idx = 5
                
                for r in rows[1:]:
                    if not r or len(r) <= max(name_idx, email_idx, cccd_idx):
                        continue
                    name_val = r[name_idx]
                    email_val = r[email_idx]
                    cccd_val = r[cccd_idx] if cccd_idx != -1 else ""
                    if name_val and email_val:
                        email_c = clean_email(email_val)
                        k_name = norm_name(name_val)
                        k_cccd = norm_cccd(cccd_val)
                        if k_name and k_cccd:
                            sso_map[(k_name, k_cccd)] = email_c
                        if k_name and k_name not in sso_map:
                            sso_map[k_name] = email_c
                        if k_cccd and k_cccd not in sso_map:
                            sso_map[k_cccd] = email_c
        except Exception as e:
            log_func(f"Lỗi đọc file Excel {filename}: {e}")

    # 2. Đọc file Word
    people_list = []
    word_files = glob.glob(os.path.join(folder_path, "*.docx"))
    total_files = len(word_files)
    
    for i, f in enumerate(word_files, start=1):
        filename = os.path.basename(f)
        if filename.startswith("~$"):
            continue
        log_func(f"\nĐang xử lý file Word ({i}/{total_files}): {filename}")
        progress_func(int((i / max(1, total_files)) * 80))
        
        try:
            doc = docx.Document(f)
            for t_idx, table in enumerate(doc.tables):
                if not table.rows:
                    continue
                header = [clean_text(cell.text) for cell in table.rows[0].cells]
                name_idx, cccd_idx, dob_idx, issue_date_idx, issue_place_idx, phone_idx, pos_idx, unit_idx = [-1]*8
                
                for idx, col in enumerate(header):
                    c_no_acc = remove_accents(col)
                    if "ho va ten" in c_no_acc or "ho ten" in c_no_acc:
                        name_idx = idx
                    elif ("cccd" in c_no_acc or "cmt" in c_no_acc or "cmnd" in c_no_acc) and "noi cap" not in c_no_acc and "ngay cap" not in c_no_acc:
                        cccd_idx = idx
                    elif "ngay sinh" in c_no_acc:
                        dob_idx = idx
                    elif "ngay cap" in c_no_acc:
                        issue_date_idx = idx
                    elif "noi cap" in c_no_acc:
                        issue_place_idx = idx
                    elif "sdt" in c_no_acc or "dien thoai" in c_no_acc:
                        phone_idx = idx
                    elif "chuc vu" in c_no_acc:
                        pos_idx = idx
                    elif "don vi" in c_no_acc:
                        unit_idx = idx
                
                if name_idx != -1 and cccd_idx != -1:
                    log_func(f"  -> Tìm thấy bảng dữ liệu (Bảng {t_idx+1})")
                    for r in table.rows[1:]:
                        vals = [clean_text(cell.text) for cell in r.cells]
                        if len(vals) > max(name_idx, cccd_idx):
                            name_val = vals[name_idx]
                            cccd_val = vals[cccd_idx]
                            if name_val and cccd_val and "họ và tên" not in name_val.lower():
                                k_name = norm_name(name_val)
                                k_cccd = norm_cccd(cccd_val)
                                email = sso_map.get((k_name, k_cccd)) or sso_map.get(k_cccd) or sso_map.get(k_name) or ""
                                dob = normalize_date(vals[dob_idx]) if dob_idx != -1 and dob_idx < len(vals) else ""
                                issue_date = normalize_date(vals[issue_date_idx]) if issue_date_idx != -1 and issue_date_idx < len(vals) else ""
                                issue_place = normalize_issue_place(vals[issue_place_idx]) if issue_place_idx != -1 and issue_place_idx < len(vals) else "CCSQLHCVTTXH"
                                phone = vals[phone_idx] if phone_idx != -1 and phone_idx < len(vals) else ""
                                pos = vals[pos_idx] if pos_idx != -1 and pos_idx < len(vals) else "Nhân viên"
                                unit = vals[unit_idx] if unit_idx != -1 and unit_idx < len(vals) else DEFAULT_ORG3
                                
                                person = {
                                    "name": name_val, "dob": dob, "cccd": format_cccd_output(cccd_val),
                                    "issue_place": issue_place, "issue_date": issue_date, "email": email,
                                    "org1": DEFAULT_ORG1, "org2": DEFAULT_ORG2, "org3": unit if unit else DEFAULT_ORG3,
                                    "org4": "", "province": DEFAULT_PROVINCE, "position": pos, "phone": phone,
                                    "serial_old": "", "network": "", "is_transfer": "", "device_type": ""
                                }
                                people_list.append(person)
                                log_func(f"     + Lấy được: {name_val} - CCCD: {format_cccd_output(cccd_val)} -> Email: {email if email else 'CHƯA CÓ SSO'}")
        except Exception as e:
            log_func(f"Lỗi đọc file Word {filename}: {e}")

    # 3. Xuất Excel
    template_file = os.path.join(folder_path, "Ket_qua.xlsx")
    if not os.path.exists(template_file):
        template_file = os.path.join(folder_path, "Template_import VGCA.xlsx")
    
    if os.path.exists(template_file):
        wb = openpyxl.load_workbook(template_file)
    else:
        wb = openpyxl.Workbook()
        
    ws = wb.active
    ws.cell(row=1, column=1, value="STT")
    ws.cell(row=1, column=2, value="Họ tên (bắt buộc)")
    ws.cell(row=1, column=3, value="Ngày sinh (bắt buộc, để dạng dd/MM/YYYY)")
    ws.cell(row=1, column=4, value="CCCC (Căn cước công dân, bắt buộc)")
    ws.cell(row=1, column=5, value="Nơi cấp (Nơi cấp CCCD; để một trong 2 nơi: CCSQLHCVTTXH hoặc CCSĐKQLCTVDLQGVDC)")
    ws.cell(row=1, column=6, value="Ngày cấp (bắt buộc, để dạng dd/MM/YYYY)")
    ws.cell(row=1, column=7, value="Địa chỉ thư điện tử công vụ (bắt buộc)")
    ws.cell(row=1, column=8, value="Tổ chức cấp 1 (bắt buộc)")
    ws.cell(row=1, column=9, value="Tổ chức cấp 2")
    ws.cell(row=1, column=10, value="Tổ chức cấp 3")
    ws.cell(row=1, column=11, value="Tổ chức cấp 4")
    ws.cell(row=1, column=12, value="Tỉnh thành phố (Bắt buộc)")
    ws.cell(row=1, column=13, value="Tỉnh/Thành phố")
    ws.cell(row=1, column=14, value="Chức vụ (Bắt buộc)")
    ws.cell(row=1, column=15, value="Điện thoại")
    ws.cell(row=1, column=16, value="Serial CTS cũ")
    ws.cell(row=1, column=17, value="Nhà mạng (bắt buộc nếu đăng ký SIM)")
    ws.cell(row=1, column=18, value="Là chuyển số (đánh dấu X nếu muốn chuyển số đang dùng sang SIM ký số)")
    ws.cell(row=1, column=19, value="Loại thiết bị( nếu không ghi sẽ là token; là một trong các loại sau: (TOKEN, SIM, SOFT, RSSP, HSM); RSSP là ký số tập trung")

    if ws.max_row > 1:
        ws.delete_rows(2, ws.max_row - 1)
        
    for idx, p in enumerate(people_list, start=1):
        row_num = idx + 1
        ws.cell(row=row_num, column=1, value=idx)
        ws.cell(row=row_num, column=2, value=p["name"])
        ws.cell(row=row_num, column=3, value=p["dob"])
        ws.cell(row=row_num, column=4, value=str(p["cccd"]))
        ws.cell(row=row_num, column=5, value=p["issue_place"])
        ws.cell(row=row_num, column=6, value=p["issue_date"])
        ws.cell(row=row_num, column=7, value=p["email"])
        ws.cell(row=row_num, column=8, value=p["org1"])
        ws.cell(row=row_num, column=9, value=p["org2"])
        ws.cell(row=row_num, column=10, value=p["org3"])
        ws.cell(row=row_num, column=11, value=p["org4"])
        ws.cell(row=row_num, column=12, value=p["province"])
        ws.cell(row=row_num, column=13, value=DEFAULT_PROVINCE)
        ws.cell(row=row_num, column=14, value=p["position"])
        ws.cell(row=row_num, column=15, value=p["phone"])
        ws.cell(row=row_num, column=16, value=p["serial_old"])
        ws.cell(row=row_num, column=17, value=p["network"])
        ws.cell(row=row_num, column=18, value=p["is_transfer"])
        ws.cell(row=row_num, column=19, value=p["device_type"])
        for col in range(1, 20):
            ws.cell(row=row_num, column=col).number_format = "@"

    output_file = os.path.join(folder_path, "Ket_qua.xlsx")
    saved = False
    for candidate in [output_file, os.path.join(folder_path, "Ket_qua_moi.xlsx"), os.path.join(folder_path, "Ket_qua_final.xlsx")]:
        try:
            wb.save(candidate)
            output_file = candidate
            saved = True
            break
        except PermissionError:
            continue
            
    if not saved:
        output_file = os.path.join(folder_path, f"Ket_qua_{int(datetime.now().timestamp())}.xlsx")
        wb.save(output_file)
        
    progress_func(100)
    log_func(f"\n==========================================")
    log_func(f"ĐÃ XUẤT THÀNH CÔNG FILE KẾT QUẢ: {os.path.basename(output_file)}")
    log_func(f"Tổng số bản ghi: {len(people_list)}")
    log_func(f"==========================================")

# ============================================================
# LOGIC TÁC VỤ 2: TỔNG HỢP XIN CẤP CKS (đọc Word + đối chiếu SSO Excel)
# ============================================================
def run_task_cks(folder_path, log_func, progress_func):
    log_func("=== BẮT ĐẦU TÁC VỤ 2: TỔNG HỢP DANH SÁCH XIN CẤP CHỮ KÝ SỐ ===")

    # ---- BƯỚC 1: Đọc file SSO Excel lấy email ---
    sso_map = {}
    excel_files = glob.glob(os.path.join(folder_path, "*.xlsx"))
    sso_found = False
    for f in excel_files:
        filename = os.path.basename(f)
        if filename.startswith("~$") or filename in ["Ket_qua.xlsx", "Ket_qua_moi.xlsx", "Ket_qua_final.xlsx", "Template_import VGCA.xlsx", "DANH_SACH_TONG_HOP.xlsx"]:
            continue
        log_func(f"Đang đọc dữ liệu SSO từ: {filename}")
        try:
            wb_sso = openpyxl.load_workbook(f, data_only=True)
            for sheetname in wb_sso.sheetnames:
                ws_sso = wb_sso[sheetname]
                rows_sso = list(ws_sso.iter_rows(values_only=True))
                if not rows_sso:
                    continue
                header_sso = [clean_text(c) for c in rows_sso[0]]
                name_idx_s, email_idx_s, cccd_idx_s = -1, -1, -1
                for idx, col in enumerate(header_sso):
                    c_lower = col.lower()
                    if ("tên" in c_lower or "họ tên" in c_lower) and name_idx_s == -1:
                        name_idx_s = idx
                    if ("tên đăng nhập" in c_lower or "email" in c_lower or "thư điện tử" in c_lower) and email_idx_s == -1:
                        email_idx_s = idx
                    if ("description" in c_lower or "cccd" in c_lower or "cmt" in c_lower or "cmnd" in c_lower) and cccd_idx_s == -1:
                        cccd_idx_s = idx
                if name_idx_s == -1: name_idx_s = 0
                if email_idx_s == -1: email_idx_s = 1
                if cccd_idx_s == -1 and len(header_sso) >= 6: cccd_idx_s = 5
                for r in rows_sso[1:]:
                    if not r or len(r) <= max(filter(lambda x: x != -1, [name_idx_s, email_idx_s, cccd_idx_s])):
                        continue
                    name_v = r[name_idx_s]
                    email_v = r[email_idx_s]
                    cccd_v = r[cccd_idx_s] if cccd_idx_s != -1 else ""
                    if name_v and email_v:
                        ec = clean_email(email_v)
                        kn = norm_name(name_v)
                        kc = norm_cccd(cccd_v)
                        if kn and kc:
                            sso_map[(kn, kc)] = ec
                        if kn and kn not in sso_map:
                            sso_map[kn] = ec
                        if kc and kc not in sso_map:
                            sso_map[kc] = ec
            sso_found = True
        except Exception as e:
            log_func(f"  Lỗi đọc file SSO {filename}: {e}")

    if sso_found:
        log_func(f"  -> Đã nạp {len(sso_map)} bản ghi SSO (email công vụ)")
    else:
        log_func("  [CẢNH BÁO] Không tìm thấy file SSO Excel trong thư mục – cột email sẽ để trống!")

    # ---- BƯỚC 2: Đọc các file Word ----
    word_files = glob.glob(os.path.join(folder_path, "*.docx"))
    total_files = len(word_files)
    records = []

    headers = [
        "STT", "Họ và tên", "Ngày sinh",
        "Số CMND/ CCCD/Hộ chiếu;Ngày cấp; nơi cấp",
        "Địa chỉ thư điện tử công vụ",
        "Tên cơ quan-tổ chức công tác",
        "Tỉnh/Thành phố", "Chức vụ",
        "Số điện thoại di động",
        "Số hiệu chứng thư cũ (nếu có)",
        "SIM PKI"
    ]

    for i, f in enumerate(word_files, start=1):
        filename = os.path.basename(f)
        if filename.startswith("~$"): continue
        log_func(f"\nĐang đọc file Word ({i}/{total_files}): {filename}")
        progress_func(int((i / max(1, total_files)) * 80))
        try:
            doc = docx.Document(f)
            for table in doc.tables:
                if not table.rows: continue
                header = [clean_text(cell.text) for cell in table.rows[0].cells]
                name_idx, dob_idx, phone_idx, cccd_idx, issue_date_idx, issue_place_idx, unit_idx, pos_idx = [-1]*8

                for idx, col in enumerate(header):
                    c_no_acc = remove_accents(col)
                    if "ho va ten" in c_no_acc or "ho ten" in c_no_acc: name_idx = idx
                    elif "ngay sinh" in c_no_acc: dob_idx = idx
                    elif "sdt" in c_no_acc or "dien thoai" in c_no_acc or "di dong" in c_no_acc: phone_idx = idx
                    elif ("cccd" in c_no_acc or "cmt" in c_no_acc or "cmnd" in c_no_acc) and "noi cap" not in c_no_acc and "ngay cap" not in c_no_acc: cccd_idx = idx
                    elif "ngay cap" in c_no_acc: issue_date_idx = idx
                    elif "noi cap" in c_no_acc: issue_place_idx = idx
                    elif "don vi" in c_no_acc or "co quan" in c_no_acc or "to chuc" in c_no_acc: unit_idx = idx
                    elif "chuc vu" in c_no_acc: pos_idx = idx

                if name_idx != -1 and cccd_idx != -1:
                    for r in table.rows[1:]:
                        vals = [clean_text(cell.text) for cell in r.cells]
                        if len(vals) > max(name_idx, cccd_idx):
                            name_val = vals[name_idx]
                            cccd_val = vals[cccd_idx]
                            if name_val and cccd_val and "họ và tên" not in name_val.lower():
                                dob_val     = normalize_date(vals[dob_idx]) if dob_idx != -1 and dob_idx < len(vals) else ""
                                phone_val   = clean_text(vals[phone_idx]) if phone_idx != -1 and phone_idx < len(vals) else ""
                                cccd_fmt    = format_cccd_output(cccd_val)
                                issue_date  = normalize_date(vals[issue_date_idx]) if issue_date_idx != -1 and issue_date_idx < len(vals) else ""
                                issue_place = normalize_issue_place(vals[issue_place_idx]) if issue_place_idx != -1 and issue_place_idx < len(vals) else "CCSQLHCVTTXH"
                                unit_val    = clean_text(vals[unit_idx]) if unit_idx != -1 and unit_idx < len(vals) else DEFAULT_ORG3
                                pos_val     = clean_text(vals[pos_idx]) if pos_idx != -1 and pos_idx < len(vals) else "Nhân viên"

                                # Đối chiếu SSO lấy email
                                k_name = norm_name(name_val)
                                k_cccd = norm_cccd(cccd_val)
                                email = sso_map.get((k_name, k_cccd)) or sso_map.get(k_cccd) or sso_map.get(k_name) or ""

                                # Cột 4: CCCD;Ngày cấp;Nơi cấp
                                cccd_str = f"{cccd_fmt};{issue_date};{issue_place}".rstrip(";")

                                records.append([name_val, dob_val, cccd_str, email, unit_val, DEFAULT_PROVINCE, pos_val, phone_val, "", ""])
                                status = f"Email: {email}" if email else "CHƯA CÓ EMAIL SSO"
                                log_func(f"  + Lấy được: {name_val} - CCCD: {cccd_fmt} -> {status}")
        except Exception as e:
            log_func(f"Lỗi đọc file {filename}: {e}")

    # ---- BƯỚC 3: Xuất file TXT ----
    output_txt = os.path.join(folder_path, "DANH_SACH_TONG_HOP.txt")
    try:
        with open(output_txt, "w", encoding="utf-8-sig", newline="") as out:
            writer = csv.writer(out, delimiter=",", quoting=csv.QUOTE_MINIMAL)
            writer.writerow(headers)
            for idx, r in enumerate(records, start=1):
                writer.writerow([idx, *r])
        log_func(f"\n==========================================")
        log_func(f"ĐÃ XUẤT THÀNH CÔNG FILE TXT CKS: DANH_SACH_TONG_HOP.txt ({len(records)} bản ghi)")
        log_func(f"==========================================")
    except Exception as e:
        log_func(f"Lỗi xuất file TXT CKS: {e}")
    progress_func(100)

# ============================================================
# LOGIC TÁC VỤ 3: TỔNG HỢP XIN CẤP EMAIL CÔNG VỤ
# ============================================================
def run_task_email(folder_path, log_func, progress_func):
    log_func("=== BẮT ĐẦU TÁC VỤ 3: TỔNG HỢP DANH SÁCH XIN CẤP EMAIL CÔNG VỤ ===")
    word_files = glob.glob(os.path.join(folder_path, "*.docx"))
    total_files = len(word_files)
    records = []
    
    headers = ["STT", "Họ và tên", "Ngày sinh", "Di động", "Số CCCD", "Đơn vị công tác", "Chức vụ", "Ghi chú"]
    
    for i, f in enumerate(word_files, start=1):
        filename = os.path.basename(f)
        if filename.startswith("~$"): continue
        log_func(f"\nĐang đọc file Word ({i}/{total_files}): {filename}")
        progress_func(int((i / max(1, total_files)) * 80))
        try:
            doc = docx.Document(f)
            for table in doc.tables:
                if not table.rows: continue
                header = [clean_text(cell.text) for cell in table.rows[0].cells]
                name_idx, dob_idx, phone_idx, cccd_idx, unit_idx, pos_idx, note_idx = [-1]*7
                
                for idx, col in enumerate(header):
                    c_no_acc = remove_accents(col)
                    if "ho va ten" in c_no_acc or "ho ten" in c_no_acc: name_idx = idx
                    elif "ngay sinh" in c_no_acc: dob_idx = idx
                    elif "sdt" in c_no_acc or "dien thoai" in c_no_acc or "di dong" in c_no_acc: phone_idx = idx
                    elif ("cccd" in c_no_acc or "cmt" in c_no_acc or "cmnd" in c_no_acc) and "noi cap" not in c_no_acc and "ngay cap" not in c_no_acc: cccd_idx = idx
                    elif "don vi" in c_no_acc or "co quan" in c_no_acc or "to chuc" in c_no_acc: unit_idx = idx
                    elif "chuc vu" in c_no_acc: pos_idx = idx
                    elif "ghi chu" in c_no_acc: note_idx = idx

                if name_idx != -1 and cccd_idx != -1:
                    for r in table.rows[1:]:
                        vals = [clean_text(cell.text) for cell in r.cells]
                        if len(vals) > max(name_idx, cccd_idx):
                            name_val = vals[name_idx]
                            cccd_val = vals[cccd_idx]
                            if name_val and cccd_val and "họ và tên" not in name_val.lower():
                                dob_val = normalize_date(vals[dob_idx]) if dob_idx != -1 and dob_idx < len(vals) else ""
                                phone_val = clean_text(vals[phone_idx]) if phone_idx != -1 and phone_idx < len(vals) else ""
                                cccd_val_fmt = format_cccd_output(cccd_val)
                                unit_val = clean_text(vals[unit_idx]) if unit_idx != -1 and unit_idx < len(vals) else DEFAULT_ORG3
                                pos_val = clean_text(vals[pos_idx]) if pos_idx != -1 and pos_idx < len(vals) else "Nhân viên"
                                note_val = clean_text(vals[note_idx]) if note_idx != -1 and note_idx < len(vals) else "Cấp mới"
                                if not note_val: note_val = "Cấp mới"
                                
                                records.append([name_val, dob_val, phone_val, cccd_val_fmt, unit_val, pos_val, note_val])
                                log_func(f"  + Lấy được Email: {name_val} - CCCD: {cccd_val_fmt} - Đơn vị: {unit_val}")
        except Exception as e:
            log_func(f"Lỗi đọc file {filename}: {e}")

    output_txt = os.path.join(folder_path, "DANH_SACH_EMAIL_CONG_VU.txt")
    try:
        with open(output_txt, "w", encoding="utf-8-sig", newline="") as out:
            writer = csv.writer(out, delimiter=",", quoting=csv.QUOTE_MINIMAL)
            writer.writerow(headers)
            for idx, r in enumerate(records, start=1):
                writer.writerow([idx, *r])
        log_func(f"\nĐÃ XUẤT THÀNH CÔNG FILE EMAIL CÔNG VỤ: DANH_SACH_EMAIL_CONG_VU.txt ({len(records)} bản ghi)")
    except Exception as e:
        log_func(f"Lỗi xuất file Email: {e}")
    progress_func(100)

# ============================================================
# GIAO DIỆN GUI TKINTER HỢP NHẤT
# ============================================================
class MainApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("PHẦN MỀM TỔNG HỢP CẤP CKS & EMAIL CÔNG VỤ (VGCA)")
        self.geometry("1000x720")
        self.minsize(900, 650)
        
        # Style
        style = ttk.Style(self)
        style.theme_use('clam')
        
        # Header Label
        header_frame = tk.Frame(self, bg="#1a365d", pady=10)
        header_frame.pack(fill=tk.X)
        lbl_title = tk.Label(
            header_frame, 
            text="HỆ THỐNG TỔNG HỢP CẤP CHỨNG THƯ SỐ & EMAIL CÔNG VỤ", 
            font=("Arial", 14, "bold"), 
            fg="white", bg="#1a365d"
        )
        lbl_title.pack()

        # Notebook Tabs
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 3 Tabs
        self.tab1 = ttk.Frame(self.notebook)
        self.tab2 = ttk.Frame(self.notebook)
        self.tab3 = ttk.Frame(self.notebook)

        self.notebook.add(self.tab1, text="  1. Đối chiếu & Xuất Ket_qua.xlsx (VGCA)  ")
        self.notebook.add(self.tab2, text="  2. Tổng hợp Cấp Chữ Ký Số (TXT/Excel)  ")
        self.notebook.add(self.tab3, text="  3. Tổng hợp Cấp Email Công Vụ  ")

        self.build_tab(self.tab1, task_type=1)
        self.build_tab(self.tab2, task_type=2)
        self.build_tab(self.tab3, task_type=3)

    def build_tab(self, tab, task_type):
        main_paned = ttk.PanedWindow(tab, orient=tk.HORIZONTAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # ----- BÊN TRÁI: HUẤN DẪN SỬ DỤNG -----
        guide_frame = ttk.LabelFrame(main_paned, text=" 📖 Hướng Dẫn Sử Dụng Bước Qua Bước ", padding=10)
        main_paned.add(guide_frame, weight=1)

        guide_text = scrolledtext.ScrolledText(guide_frame, wrap=tk.WORD, font=("Segoe UI", 10), bg="#f8fafc", width=38)
        guide_text.pack(fill=tk.BOTH, expand=True)

        if task_type == 1:
            guide_content = """🎯 TÁC VỤ 1: ĐỐI CHIẾU WORD & EXCEL SSO (VGCA)

📌 Mục đích:
Tự động đọc thông tin từ các đơn đăng ký Word (.docx), ghép đôi (Họ tên + CCCD) với file tài khoản SSO Excel để lấy Email công vụ, tạo file Ket_qua.xlsx đúng 19 cột chuẩn VGCA.

📝 Các bước thực hiện:
1️⃣ Bước 1: Chuẩn bị file
   - Đặt các file Word (.docx) đề nghị cấp CKS vào 1 thư mục.
   - Đặt file chứa tài khoản SSO (ví dụ: Copy of SSO - BVĐKso2.xlsx) vào cùng thư mục đó.
2️⃣ Bước 2: Chọn thư mục
   - Bấm nút 'Chọn Thư Mục Dữ Liệu' và chọn thư mục chứa các file trên.
3️⃣ Bước 3: Thực hiện
   - Bấm nút 'BẮT ĐẦU XỬ LÝ ĐỐI CHIẾU'.
4️⃣ Bước 4: Kiểm tra kết quả
   - Kết quả xuất ra file Ket_qua.xlsx (hoặc Ket_qua_moi.xlsx nếu file cũ đang mở) có đầy đủ 19 cột.
"""
        elif task_type == 2:
            guide_content = """🎯 TÁC VỤ 2: TỔNG HỢP DANH SÁCH XIN CẤP CKS

📌 Mục đích:
Quét tất cả file Word (.docx) trong thư mục để tổng hợp danh sách xin cấp mới/cấp lại Chữ ký số ra file TXT danh sách tổng hợp.

📝 Các bước thực hiện:
1️⃣ Bước 1: Đặt các file Word mẫu đơn CKS vào thư mục cần xử lý.
2️⃣ Bước 2: Bấm nút 'Chọn Thư Mục Dữ Liệu' và chọn thư mục chứa file.
3️⃣ Bước 3: Bấm nút 'BẮT ĐẦU TỔNG HỢP CKS'.
4️⃣ Bước 4: Kiểm tra file kết quả DANH_SACH_TONG_HOP.txt được tự động tạo trong thư mục.
"""
        else:
            guide_content = """🎯 TÁC VỤ 3: TỔNG HỢP DANH SÁCH XIN CẤP EMAIL CÔNG VỤ

📌 Mục đích:
Quét tất cả file Word (.docx) trong thư mục để tổng hợp danh sách đề nghị cấp tài khoản Thư điện tử công vụ mới.

📝 Các bước thực hiện:
1️⃣ Bước 1: Đặt các file Word mẫu xin cấp Email công vụ vào thư mục.
2️⃣ Bước 2: Bấm nút 'Chọn Thư Mục Dữ Liệu' và chọn thư mục chứa file.
3️⃣ Bước 3: Bấm nút 'BẮT ĐẦU TỔNG HỢP EMAIL'.
4️⃣ Bước 4: Kiểm tra file kết quả DANH_SACH_EMAIL_CONG_VU.txt được tự động tạo trong thư mục.
"""
        guide_text.insert(tk.END, guide_content)
        guide_text.config(state=tk.DISABLED)

        # ----- BÊN PHẢI: KHU VỰC THỰC THI & LOG -----
        action_frame = ttk.LabelFrame(main_paned, text=" ⚡ Thao Tác Thực Thi ", padding=10)
        main_paned.add(action_frame, weight=2)

        # Chọn thư mục
        folder_frame = ttk.Frame(action_frame)
        folder_frame.pack(fill=tk.X, pady=5)

        ttk.Label(folder_frame, text="Thư mục làm việc:", font=("Segoe UI", 10, "bold")).pack(side=tk.LEFT, padx=5)
        txt_folder = ttk.Entry(folder_frame, font=("Segoe UI", 10))
        txt_folder.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        txt_folder.insert(0, os.getcwd())

        def browse_folder():
            selected = filedialog.askdirectory(initialdir=txt_folder.get())
            if selected:
                txt_folder.delete(0, tk.END)
                txt_folder.insert(0, selected)

        ttk.Button(folder_frame, text="📁 Chọn Thư Mục", command=browse_folder).pack(side=tk.RIGHT, padx=5)

        # Thanh tiến trình (Progress Bar)
        progress_var = tk.DoubleVar()
        progress_bar = ttk.Progressbar(action_frame, variable=progress_var, maximum=100)
        progress_bar.pack(fill=tk.X, pady=10, padx=5)

        # Khung nhật ký xử lý (Log Console)
        log_box = scrolledtext.ScrolledText(action_frame, height=16, font=("Consolas", 10), bg="#1e293b", fg="#e2e8f0")
        log_box.pack(fill=tk.BOTH, expand=True, pady=5, padx=5)

        def append_log(msg):
            log_box.insert(tk.END, msg + "\n")
            log_box.see(tk.END)

        def update_progress(val):
            progress_var.set(val)
            self.update_idletasks()

        # Nút Chạy tác vụ
        def start_worker():
            folder = txt_folder.get().strip()
            if not os.path.exists(folder):
                messagebox.showerror("Lỗi", "Thư mục chọn không tồn tại!")
                return
            
            btn_run.config(state=tk.DISABLED)
            log_box.delete("1.0", tk.END)
            progress_var.set(0)

            def thread_target():
                try:
                    if task_type == 1:
                        run_task_doi_chieu(folder, append_log, update_progress)
                    elif task_type == 2:
                        run_task_cks(folder, append_log, update_progress)
                    else:
                        run_task_email(folder, append_log, update_progress)
                    messagebox.showinfo("Thành công", "Tác vụ đã hoàn thành thành công!")
                except Exception as ex:
                    append_log(f"\n❌ LỖI HỆ THỐNG: {ex}")
                    messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {ex}")
                finally:
                    btn_run.config(state=tk.NORMAL)

            threading.Thread(target=thread_target, daemon=True).start()

        btn_text = "🚀 BẮT ĐẦU XỬ LÝ ĐỐI CHIẾU (VGCA)" if task_type == 1 else ("🚀 BẮT ĐẦU TỔNG HỢP CKS" if task_type == 2 else "🚀 BẮT ĐẦU TỔNG HỢP EMAIL")
        btn_run = tk.Button(
            action_frame, text=btn_text, font=("Segoe UI", 11, "bold"), 
            bg="#2563eb", fg="white", activebackground="#1d4ed8", activeforeground="white",
            pady=8, command=start_worker
        )
        btn_run.pack(fill=tk.X, pady=10, padx=5)

if __name__ == "__main__":
    app = MainApp()
    app.mainloop()
